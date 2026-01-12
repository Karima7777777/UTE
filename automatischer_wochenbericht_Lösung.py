# Importiere die Bibliotheken pandas, matplotlib.pyplot und docx - importiere das aktuelle Datum
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document # für Word-Dokumente
from docx.shared import Inches, Cm
from datetime import date # für das aktuelle Datum

# --- DATENANALYSE ---

# Lade die CSV-Datei 'wochen_verkaufsdaten.csv' und konvertiere die 'Order Date'-Spalte in ein Datumsformat. Setze diese Spalte dann als Index der Tabelle.
df = pd.read_csv('wochen_verkaufsdaten.csv', parse_dates=['Order Date'], index_col='Order Date')


# Aggregiere die 'Sales'-Spalte auf wöchentlicher Basis (W) und bilde die Summe. Speichere das Ergebnis in 'woechentliche_verkaeufe'.
woechentliche_verkaeufe = df.resample('W')['Sales'].sum()


# --- GRAFIK ERSTELLEN UND SPEICHERN ---

# Erstelle eine neue Grafik mit einer angenehmen Größe
plt.figure(figsize=(10, 5))

# Erstelle ein Liniendiagramm für die woechentlichen_verkaeufe
plt.plot(woechentliche_verkaeufe.index, woechentliche_verkaeufe.values, marker='o', linestyle='-', color='b')

# Gib der Grafik den Titel 'Wöchentliche Verkaufszahlen'
plt.title('Wöchentliche Verkaufszahlen')

# Beschrifte die y-Achse mit 'Umsatz in USD'
plt.ylabel('Umsatz in USD')
# Beschrifte die x-Achse mit 'Woche'
plt.xlabel('Woche')

# Speichere die Grafik in einer Datei namens 'wochenbericht_grafik.png' mit hoher Auflösung
plt.savefig('wochenbericht_grafik.png', dpi=300)

# --- BERICHTS-INFORMATIONEN VORBEREITEN ---

# 1. Die Uhr ablesen
heute = date.today()

# 2. Das Datum in zwei Formaten aufbereiten
datum_fuer_bericht = heute.strftime('%d.%m.%Y') # Format "Tag.Monat.Jahr"
datum_fuer_dateiname = heute.strftime('%Y-%m-%d') # Format "Jahr-Monat-Tag" für saubere Dateinamen


# --- KI-ZUSAMMENFASSUNG ERSTELLEN ---

# Erstelle eine kurze, professionelle Zusammenfassung in einem Satz für Manager.
zusammenfassung = "Die Daten zeigen, dass die wöchentlichen Verkäufe nach einem starken Start schwanken, aber insgesamt einen leichten Aufwärtstrend aufweisen."
ki_zusammenfassung = zusammenfassung


# --- WORD-BERICHT ERSTELLEN ---

# Importiere die Bibliothek für die Arbeit mit Word-Dokumenten
from docx import Document


# Öffne das Word-Dokument 'berichtsvorlage.docx'
doc = Document('berichtsvorlage.docx')

# Führe "Suchen & Ersetzen" für das gesamte Dokument aus
for p in doc.paragraphs:
    # Ersetze Text-Platzhalter
    p.text = p.text.replace('{{TITEL}}', 'Wöchentlicher Verkaufsbericht')
    p.text = p.text.replace('{{ZUSAMMENFASSUNG}}', ki_zusammenfassung)
    p.text = p.text.replace('{{Date}}', datum_fuer_bericht)
    
    # Ersetze den Grafik-Platzhalter
    if '{{GRAFIK}}' in p.text:
        p.text = '' # Leere den Absatz
        p.add_run().add_picture('wochenbericht_grafik.png', width=Cm(15.0))


# --- SPEICHERE DAS DOKUMENT ---
# Speichere das fertige Dokument unter dem Namen 'finaler_wochenbericht.docx'
# doc.save('finaler_wochenbericht.docx')

# Erstelle den dynamischen Dateinamen
finaler_dateiname = f"wochenbericht_{datum_fuer_dateiname}.docx"

# Speichere das fertige Dokument unter dem neuen, dynamischen Namen
doc.save(finaler_dateiname)

print("Der automatisierte Word-Bericht wurde erfolgreich erstellt!")